import os
from typing import List, Union
from langchain.document_loaders import UnstructuredWordDocumentLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
import pandas as pd
from langchain_openai import ChatOpenAI
from docx import Document
import sys
import os
from typing import Dict, List, Tuple, Any
os.environ['KMP_DUPLICATE_LIB_OK'] = 'TRUE'
import asyncio

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))
from Input_extraction.extract_polarion_field_mapping import *
from utils.simple_functions import *
from llm.llm import a_invoke_model
from Processing.copertura_requisiti import add_new_TC, save_updated_json
from Processing.controllo_sintattico import fill_excel_file
from Processing.test_design import process_paragraphs, process_docx

embedding_model = "text-embedding-3-large"
PANDOC_EXE = "pandoc" 


def process_docx_with_diagnostics(docx_path: str, output_dir: str) -> List[str]:
    """
    Process DOCX with comprehensive diagnostics to identify issues
    """
    print(f"\n=== DIAGNOSTIC INFO ===")
    print(f"Input path: {docx_path}")
    print(f"File exists: {os.path.exists(docx_path)}")
    print(f"Output dir: {output_dir}")
    
    try:
        # Your existing process_docx logic here
        # This is where the conversion happens
        result = process_docx(docx_path, output_dir)
        
        # CRITICAL DIAGNOSTIC: Check what process_docx returns
        print(f"\n=== PROCESS_DOCX OUTPUT ===")
        print(f"Type of result: {type(result)}")
        print(f"Length: {len(result) if hasattr(result, '__len__') else 'N/A'}")
        
        if result and len(result) > 0:
            print(f"First item type: {type(result[0])}")
            print(f"First item has page_content: {hasattr(result[0], 'page_content')}")
            
            # Check if it's a list of strings or Document objects
            if isinstance(result[0], str):
                print(f"First item (string): {result[0][:100]}...")
            elif hasattr(result[0], 'page_content'):
                print(f"First item page_content: {result[0].page_content[:100]}...")
            else:
                print(f"First item: {str(result[0])[:100]}...")
        
        return result
        
    except Exception as e:
        print(f"[ERROR] process_docx failed: {e}")
        import traceback
        traceback.print_exc()
        raise


def safe_extract_text(chunks) -> List[str]:
    """
    Safely extract text from any chunk format.
    Handles: strings, Document objects, lists, or any other format
    """
    text_chunks = []
    
    print(f"\n=== EXTRACTING TEXT ===")
    print(f"Input type: {type(chunks)}")
    print(f"Input length: {len(chunks) if hasattr(chunks, '__len__') else 'N/A'}")
    
    # If chunks is not a list, make it one
    if not isinstance(chunks, list):
        chunks = [chunks]
    
    for i, chunk in enumerate(chunks):
        try:
            # Case 1: Document object with page_content
            if hasattr(chunk, 'page_content'):
                text = chunk.page_content
                if i == 0:
                    print(f"Chunk format: Document with page_content")
            
            # Case 2: Already a string
            elif isinstance(chunk, str):
                text = chunk
                if i == 0:
                    print(f"Chunk format: String")
            
            # Case 3: Dictionary with text/content field
            elif isinstance(chunk, dict):
                text = chunk.get('text') or chunk.get('content') or chunk.get('page_content') or str(chunk)
                if i == 0:
                    print(f"Chunk format: Dictionary")
            
            # Case 4: List (nested structure)
            elif isinstance(chunk, list):
                # Recursively process nested lists
                text = ' '.join(safe_extract_text(chunk))
                if i == 0:
                    print(f"Chunk format: Nested list")
            
            # Case 5: Other types - convert to string
            else:
                text = str(chunk)
                if i == 0:
                    print(f"Chunk format: {type(chunk).__name__} (converted to string)")
            
            # Clean and validate text
            if text and isinstance(text, str) and text.strip():
                text_chunks.append(text.strip())
        
        except Exception as e:
            print(f"[WARNING] Failed to extract text from chunk {i}: {e}")
            continue
    
    print(f"Successfully extracted {len(text_chunks)} text chunks")
    if text_chunks:
        print(f"First chunk preview: {text_chunks[0][:100]}...")
    
    return text_chunks



def safe_create_vectordb(paragraph, vectorstore, k=3, similarity_threshold=0.75):
    """
    Create vector DB query with proper text extraction
    """
    # Extract text from paragraph
    if hasattr(paragraph, 'page_content'):
        query_text = paragraph.page_content
    elif isinstance(paragraph, str):
        query_text = paragraph
    elif isinstance(paragraph, dict):
        query_text = paragraph.get('text') or paragraph.get('content') or str(paragraph)
    else:
        query_text = str(paragraph)
    
    # Validate query text
    if not query_text or not isinstance(query_text, str):
        print(f"[WARNING] Invalid query text: {type(query_text)}")
        return None
    
    try:
        docs_found = vectorstore.similarity_search_with_score(query_text, k)
        
        closest_test = []
        if docs_found:
            closest_doc, score = docs_found[0]
            score = 1 - (score / 2)
            if score >= similarity_threshold:  
                print(f"Score: {score}")
                closest_test.append(closest_doc.page_content)
            else:
                closest_test = None
        
        return closest_test
    
    except Exception as e:
        print(f"[ERROR] Vector search failed: {e}")
        print(f"Query text type: {type(query_text)}")
        print(f"Query text preview: {str(query_text)[:100]}")
        raise


async def main():
    """
    Updated main function with diagnostics
    """
    # File paths
    input_path = os.path.join(
        os.path.dirname(__file__), "..", "input",
        "2ESEMPI Requirement", "2ESEMPI Requirement",
        "PRJ0015372 - ZENIT Phase 1 - FA - Rev 1.0.docx"
    )
    
    print(f"\n=== PROCESSING INPUT FILE ===")
    print(f"Input path: {input_path}")
    print(f"File exists: {os.path.exists(input_path)}")
    
    # Process paragraphs
    paragraphs = process_docx_with_diagnostics(input_path, os.path.dirname(input_path))
    print(f"Paragraphs loaded: {len(paragraphs)}")
    
    # Process RAG chunks
    rag_path = os.path.join(
        os.path.dirname(__file__), "..", "input",
        "2ESEMPI Requirement", "2ESEMPI Requirement",
        "RU_ZENIT_V_0.4_FASE_1 (1).docx"
    )
    
    print(f"\n=== PROCESSING RAG FILE ===")
    print(f"RAG path: {rag_path}")
    print(f"File exists: {os.path.exists(rag_path)}")
    
    chunks = process_docx_with_diagnostics(input_path, os.path.dirname(rag_path))
    
    # CRITICAL FIX: Extract text safely
    text_chunks = safe_extract_text(chunks)
    
    if not text_chunks:
        raise ValueError("No valid text chunks extracted from document!")
    
    # Create vector store
    print(f"\n=== CREATING VECTOR STORE ===")
    embeddings = OpenAIEmbeddings(model=embedding_model)
    vectorstore = FAISS.from_texts(text_chunks, embeddings)
    print("Vector store created successfully!")
    
    # Rest of your code...
    mapping = extract_field_mapping()
    print("Finishing mapping")
    
    # Replace create_vectordb with safe_create_vectordb in process_paragraphs
    new_TC = await process_paragraphs(paragraphs, vectorstore, mapping)
    
    # Continue with rest of processing...


# Quick test function to compare files
def compare_docx_files(file1_path: str, file2_path: str):
    """
    Compare two DOCX files to see what's different
    """
    print(f"\n=== COMPARING FILES ===")
    
    for path in [file1_path, file2_path]:
        print(f"\nFile: {os.path.basename(path)}")
        print(f"Exists: {os.path.exists(path)}")
        if os.path.exists(path):
            print(f"Size: {os.path.getsize(path)} bytes")
            
            # Try to check structure
            try:
                import docx
                doc = docx.Document(path)
                print(f"Paragraphs: {len(doc.paragraphs)}")
                print(f"Tables: {len(doc.tables)}")
                print(f"Has images: {len(doc.inline_shapes) > 0}")
            except Exception as e:
                print(f"Could not analyze with python-docx: {e}")


if __name__ == "__main__":
    # First, compare the files to see what's different
    file1 = "path/to/RU_ZENIT_V_0.4_FASE_1.docx"  # Working file
    file2 = "path/to/PRJ0015372 - ZENIT Phase 1 - FA - Rev 1.0.docx"  # Problematic file
    
    # Uncomment to compare:
    # compare_docx_files(file1, file2)
    
    asyncio.run(main())